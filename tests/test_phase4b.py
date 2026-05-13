"""Phase 4B ship test — parsers + DocumentChunker.

Phase 4B introduces the path from `bytes` on disk to `DocChunk`s ready
for embedding. The data layer is already proven by 4A; this test
proves the *content* layer:

  1. PDF parser: synthesizes a 3-page PDF on the fly, asserts one block
     per non-empty page, page_number is 1-indexed.
  2. DOCX parser: synthesizes a doc with H1 -> H2 -> paragraph -> table
     and asserts section_path lineage + table row flattening.
  3. XLSX parser: synthesizes a workbook with a header row and a
     no-header sheet; asserts header-attaching format vs col_A fallback.
  4. Dispatcher: picks subtype from mime_type then extension; raises
     `UnsupportedDocumentError` for octet-stream with no recognisable ext.
  5. DocumentChunker: handed a sequence of `ParsedBlock`s with mixed
     page/section, produces chunks that
        - never exceed target_tokens (modulo over-budget single blocks)
        - inherit `page_number` and `section_path` from the dominant block
        - apply 100-token overlap (chunk N+1 head shares text with chunk N tail)
        - emit `pages_covered` when a chunk straddles a page boundary

Run with:

    venv\\Scripts\\python.exe tests\\test_phase4b.py
"""
from __future__ import annotations

import io
import os
import sys
import traceback
from contextlib import contextmanager
from typing import Callable, List, Tuple

_HERE = os.path.dirname(os.path.abspath(__file__))
_ROOT = os.path.dirname(_HERE)
if _ROOT not in sys.path:
    sys.path.insert(0, _ROOT)


results: List[Tuple[str, str, str, str]] = []


@contextmanager
def section(label: str):
    print(f"\n=== {label} ===")
    yield


def check(slice_id: str, name: str, fn: Callable[[], None]) -> None:
    try:
        fn()
    except AssertionError as e:
        msg = str(e) or "assertion failed"
        results.append((slice_id, name, "FAIL", msg))
        print(f"  [FAIL] {name} :: {msg}")
        return
    except Exception:
        msg = traceback.format_exc(limit=4).strip().splitlines()[-1]
        results.append((slice_id, name, "FAIL", msg))
        print(f"  [ERROR] {name} :: {msg}")
        return
    results.append((slice_id, name, "PASS", ""))
    print(f"  [PASS] {name}")


# ---------------------------------------------------------------------------
# Synthetic doc builders
# ---------------------------------------------------------------------------

def _build_pdf_three_pages() -> bytes:
    """Build a minimal 3-page PDF with one ASCII text block per page.
    Hand-rolled because reportlab is heavy and we only need text-extract
    round-trip, not formatting fidelity. pypdf accepts this happily."""
    pages_text = [
        "Page one introduces the project Helios with deliverables.",
        "Page two covers Q3 milestones and dependencies on team Hydra.",
        "Page three lists risks and mitigations for the Helios launch.",
    ]
    # Build each page content stream
    pdf = io.BytesIO()
    pdf.write(b"%PDF-1.4\n")
    objects: list[bytes] = []
    offsets: list[int] = []

    def add_obj(data: bytes) -> int:
        offsets.append(pdf.tell())
        obj_id = len(objects) + 1
        pdf.write(f"{obj_id} 0 obj\n".encode("ascii"))
        pdf.write(data)
        pdf.write(b"\nendobj\n")
        objects.append(data)
        return obj_id

    # Catalog (1), Pages (2), Font (3), then page+content pairs
    catalog_id = add_obj(b"<< /Type /Catalog /Pages 2 0 R >>")
    # Reserve pages object — emit after we know kids
    pages_offset_idx = len(offsets)
    offsets.append(0)  # placeholder
    pdf.write(b"2 0 obj\n")
    pages_obj_pos = pdf.tell()
    pdf.write(b"<< /Type /Pages /Count 3 /Kids [")
    # We'll patch kids below
    placeholder_pos = pdf.tell()
    pdf.write(b"                                   ")  # 35 bytes placeholder
    pdf.write(b"] /MediaBox [0 0 612 792] >>\nendobj\n")
    objects.append(b"placeholder")  # for indexing

    # Update offsets[1] now
    offsets[pages_offset_idx] = offsets[pages_offset_idx] or 0
    # Actually overwrite the offsets entry for object 2:
    # we tracked add_obj's offset only via list append; we wrote object 2 manually
    # so we need to set offsets[1] manually:
    # find current pages obj start: it's right after object 1.
    # We'll just compute it correctly:
    # Easier path: rebuild PDF using a simpler layout.

    return _build_pdf_three_pages_simple(pages_text)


def _build_pdf_three_pages_simple(pages_text: list[str]) -> bytes:
    """Use pypdf's PdfWriter to round-trip pages. We pull empty pages
    from a blank document and inject text via the high-level API."""
    from pypdf import PdfWriter
    from pypdf.generic import (
        DictionaryObject, NameObject, NumberObject, ArrayObject,
        IndirectObject, TextStringObject, ContentStream,
    )

    writer = PdfWriter()

    # pypdf 6.x has no direct "add page with text" — easiest is to use
    # reportlab if available. Fallback: build a hand-crafted PDF.
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError:
        return _build_pdf_three_pages_handcrafted(pages_text)

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    for text in pages_text:
        c.setFont("Helvetica", 12)
        c.drawString(72, 720, text)
        c.showPage()
    c.save()
    return buf.getvalue()


def _build_pdf_three_pages_handcrafted(pages_text: list[str]) -> bytes:
    """Build a minimal valid PDF with three pages, each containing one
    text-show operator. Compatible with pypdf's text extractor.

    Structure:
        1: Catalog
        2: Pages (Kids -> [3,5,7], Count 3, MediaBox)
        3,5,7: Page objects (Parent -> 2, Contents -> 4,6,8, Resources -> Font)
        4,6,8: Content streams
        9: Font (Helvetica)
    """

    def text_to_stream(text: str) -> bytes:
        # PDF strings escape ( ) \
        esc = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        body = (
            f"BT\n/F1 12 Tf\n72 720 Td\n({esc}) Tj\nET\n"
        ).encode("ascii")
        return body

    parts: list[bytes] = []
    offsets: list[int] = []
    out = io.BytesIO()
    out.write(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")

    def write_obj(obj_id: int, body: bytes):
        offsets.append(out.tell())
        out.write(f"{obj_id} 0 obj\n".encode("ascii"))
        out.write(body)
        out.write(b"\nendobj\n")

    # 1: Catalog
    write_obj(1, b"<< /Type /Catalog /Pages 2 0 R >>")
    # 2: Pages (placeholder, we'll come back to it after computing kids)
    # We can write it now since IDs are known:
    write_obj(2, b"<< /Type /Pages /Count 3 /Kids [3 0 R 5 0 R 7 0 R] "
                 b"/MediaBox [0 0 612 792] >>")

    # Pages + contents interleaved: 3,4 then 5,6 then 7,8
    next_id = 3
    for text in pages_text:
        page_id = next_id
        content_id = next_id + 1
        # Page
        page_body = (
            f"<< /Type /Page /Parent 2 0 R "
            f"/Resources << /Font << /F1 9 0 R >> >> "
            f"/Contents {content_id} 0 R >>"
        ).encode("ascii")
        write_obj(page_id, page_body)
        # Content stream
        stream = text_to_stream(text)
        content_body = (
            f"<< /Length {len(stream)} >>\nstream\n".encode("ascii")
            + stream
            + b"endstream"
        )
        write_obj(content_id, content_body)
        next_id += 2

    # 9: Font
    write_obj(9, b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    xref_offset = out.tell()
    out.write(f"xref\n0 {10}\n".encode("ascii"))
    out.write(b"0000000000 65535 f \n")
    for off in offsets:
        out.write(f"{off:010d} 00000 n \n".encode("ascii"))
    out.write(b"trailer\n")
    out.write(b"<< /Size 10 /Root 1 0 R >>\n")
    out.write(b"startxref\n")
    out.write(f"{xref_offset}\n".encode("ascii"))
    out.write(b"%%EOF\n")
    return out.getvalue()


def _build_docx() -> bytes:
    """H1 -> H2 -> paragraph -> table.

    Section stack lineage we expect:
        H1 "Project Helios"               -> emitted with section_path=None
        H2 "Q3 Milestones"                -> emitted with section_path="Project Helios"
        para "Ship encrypted backups..."  -> emitted with section_path="Project Helios / Q3 Milestones"
        table row "Team | Owner"          -> section_path="Project Helios / Q3 Milestones"
    """
    from docx import Document
    doc = Document()
    doc.add_heading("Project Helios", level=1)
    doc.add_heading("Q3 Milestones", level=2)
    doc.add_paragraph("Ship encrypted backups by end of quarter to the EU region.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Team"
    table.rows[0].cells[1].text = "Owner"
    table.rows[1].cells[0].text = "Hydra"
    table.rows[1].cells[1].text = "Alice"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_xlsx() -> bytes:
    """Sheet 'Forecast' has a header row; sheet 'Notes' has none."""
    from openpyxl import Workbook
    wb = Workbook()
    ws1 = wb.active
    ws1.title = "Forecast"
    ws1.append(["Region", "Q3 Revenue", "Owner"])
    ws1.append(["EU", 1200000, "Alice"])
    ws1.append(["US", 1500000, "Bob"])

    ws2 = wb.create_sheet("Notes")
    ws2.append(["raw notes 1", 42])
    ws2.append(["raw notes 2", 99])

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Parser tests
# ---------------------------------------------------------------------------

def test_pdf_parser_pages():
    from app.parsers import parse_document
    pdf_bytes = _build_pdf_three_pages_simple([
        "Page one introduces the project Helios with deliverables.",
        "Page two covers Q3 milestones and dependencies on team Hydra.",
        "Page three lists risks and mitigations for the Helios launch.",
    ])
    subtype, blocks = parse_document(pdf_bytes, "application/pdf", "helios.pdf")
    assert subtype == "pdf"
    assert len(blocks) == 3, f"expected 3 blocks, got {len(blocks)} :: {[b.text[:40] for b in blocks]}"
    for i, b in enumerate(blocks, start=1):
        assert b.page_number == i, f"block {i} page_number={b.page_number}"
        assert b.section_path is None
        assert b.metadata.get("source_subtype") == "pdf"
    assert "Helios" in blocks[0].text
    assert "Hydra" in blocks[1].text


def test_docx_parser_section_path():
    from app.parsers import parse_document
    subtype, blocks = parse_document(
        _build_docx(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "helios.docx",
    )
    assert subtype == "docx"
    by_text = {b.text: b for b in blocks}

    # H1: section_path should be None (it lives at the top)
    h1 = by_text["Project Helios"]
    assert h1.section_path is None
    assert h1.metadata.get("block_role") == "heading"
    assert h1.metadata.get("heading_level") == 1

    # H2: should be under H1
    h2 = by_text["Q3 Milestones"]
    assert h2.section_path == "Project Helios", f"H2 section_path={h2.section_path!r}"
    assert h2.metadata.get("heading_level") == 2

    # Paragraph: under H1 / H2
    para = by_text["Ship encrypted backups by end of quarter to the EU region."]
    assert para.section_path == "Project Helios / Q3 Milestones"
    assert para.metadata.get("block_role") == "paragraph"

    # Table: rows flattened with " | "
    table_rows = [b for b in blocks if b.metadata.get("block_role") == "table_row"]
    assert len(table_rows) == 2, f"expected 2 table rows, got {len(table_rows)}"
    assert table_rows[0].text == "Team | Owner"
    assert table_rows[1].text == "Hydra | Alice"
    for tr in table_rows:
        assert tr.section_path == "Project Helios / Q3 Milestones"


def test_xlsx_parser_header_attach_and_fallback():
    from app.parsers import parse_document
    subtype, blocks = parse_document(
        _build_xlsx(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "forecast.xlsx",
    )
    assert subtype == "xlsx"

    forecast = [b for b in blocks if b.section_path == "Forecast"]
    notes = [b for b in blocks if b.section_path == "Notes"]
    assert len(forecast) == 2, f"expected 2 Forecast rows, got {len(forecast)}"
    assert len(notes) == 2, f"expected 2 Notes rows, got {len(notes)}"

    # Header attached to Forecast rows.
    assert "Region: EU" in forecast[0].text
    assert "Q3 Revenue: 1200000" in forecast[0].text
    assert "Owner: Alice" in forecast[0].text

    # Notes has no header row -> col_A / col_B fallback.
    assert "col_A: raw notes 1" in notes[0].text
    assert "col_B: 42" in notes[0].text


def test_dispatcher_mime_and_extension_fallback():
    from app.parsers import parse_document, UnsupportedDocumentError

    # Mime mismatch but valid extension -> still works.
    subtype, blocks = parse_document(
        _build_xlsx(), "application/octet-stream", "forecast.xlsx",
    )
    assert subtype == "xlsx" and len(blocks) > 0

    # No mime + no extension -> raises.
    try:
        parse_document(b"hello world", None, "no_extension_here")
    except UnsupportedDocumentError:
        pass
    else:
        raise AssertionError("expected UnsupportedDocumentError")


# ---------------------------------------------------------------------------
# Chunker tests
# ---------------------------------------------------------------------------

def test_chunker_respects_target_tokens():
    from app.parsers import ParsedBlock
    from app.services.document_chunker import DocumentChunker

    # 5 small blocks across 2 pages of one section.
    blocks = [
        ParsedBlock(0, "First block of text in section alpha.", page_number=1,
                    section_path="Alpha", metadata={"source_subtype": "pdf"}),
        ParsedBlock(1, "Second block continues section alpha.", page_number=1,
                    section_path="Alpha", metadata={"source_subtype": "pdf"}),
        ParsedBlock(2, "Third block bridges into page two.", page_number=2,
                    section_path="Alpha", metadata={"source_subtype": "pdf"}),
        ParsedBlock(3, "Fourth block keeps going on page two.", page_number=2,
                    section_path="Alpha", metadata={"source_subtype": "pdf"}),
        ParsedBlock(4, "Fifth block closes out the document.", page_number=2,
                    section_path="Alpha", metadata={"source_subtype": "pdf"}),
    ]
    chunker = DocumentChunker(target_tokens=40, overlap_tokens=8)
    chunks = chunker.chunk(blocks)
    assert chunks, "expected at least one chunk"
    for c in chunks:
        # Allow last chunk to overshoot via overlap-only insertion never
        # happens, but a single block can exceed budget — assert per-chunk
        # is at most target + a fudge for tokenizer rounding.
        assert c.token_count <= 40 + 32, f"chunk {c.chunk_index} = {c.token_count} tokens"
        assert c.section_path == "Alpha"
        assert c.page_number in (1, 2)


def test_chunker_overlap_preserves_continuity():
    from app.parsers import ParsedBlock
    from app.services.document_chunker import DocumentChunker

    long_text = " ".join([f"sentence number {i} alpha bravo charlie." for i in range(50)])
    blocks = [ParsedBlock(0, long_text, page_number=1, section_path="Sec", metadata={"source_subtype": "pdf"})]

    chunker = DocumentChunker(target_tokens=60, overlap_tokens=15)
    chunks = chunker.chunk(blocks)
    assert len(chunks) >= 2, f"expected multi-chunk split, got {len(chunks)}"

    # Overlap check: the head of chunk N+1 should share at least one
    # sentence (or partial sentence) with the tail of chunk N.
    for i in range(len(chunks) - 1):
        tail = chunks[i].text[-80:]
        head = chunks[i + 1].text[:80]
        # Find any 4-word shingle in common.
        tail_words = tail.split()
        head_words = head.split()
        shingles_tail = {" ".join(tail_words[j : j + 4]) for j in range(len(tail_words) - 3)}
        shingles_head = {" ".join(head_words[j : j + 4]) for j in range(len(head_words) - 3)}
        assert shingles_tail & shingles_head, (
            f"no overlap between chunk {i} tail and chunk {i+1} head"
        )


def test_chunker_dominant_page_inheritance():
    """When a chunk spans pages 1+2 with page 1 contributing more tokens,
    the chunk's page_number is 1 and `pages_covered` lists [1, 2]."""
    from app.parsers import ParsedBlock
    from app.services.document_chunker import DocumentChunker

    big_p1 = "alpha " * 30 + "."  # ~30 tokens on page 1
    small_p2 = "beta gamma."           # tiny on page 2
    blocks = [
        ParsedBlock(0, big_p1, page_number=1, section_path="Sec",
                    metadata={"source_subtype": "pdf"}),
        ParsedBlock(1, small_p2, page_number=2, section_path="Sec",
                    metadata={"source_subtype": "pdf"}),
    ]
    chunker = DocumentChunker(target_tokens=500, overlap_tokens=20)
    chunks = chunker.chunk(blocks)
    assert len(chunks) == 1
    c = chunks[0]
    assert c.page_number == 1, f"dominant page should be 1, got {c.page_number}"
    assert c.metadata.get("pages_covered") == [1, 2]


def test_chunker_empty_and_whitespace_blocks_are_skipped():
    from app.parsers import ParsedBlock
    from app.services.document_chunker import DocumentChunker

    blocks = [
        ParsedBlock(0, "  ", page_number=1, metadata={}),
        ParsedBlock(1, "", page_number=1, metadata={}),
        ParsedBlock(2, "real content here", page_number=1, metadata={"source_subtype": "pdf"}),
    ]
    chunker = DocumentChunker(target_tokens=200, overlap_tokens=10)
    chunks = chunker.chunk(blocks)
    assert len(chunks) == 1
    assert "real content here" in chunks[0].text


def test_chunker_oversize_single_block_falls_back():
    """A single block bigger than `target_tokens` should still produce
    multiple chunks via sentence splitting + hard window."""
    from app.parsers import ParsedBlock
    from app.services.document_chunker import DocumentChunker

    long = ". ".join(f"sentence {i} with several filler words" for i in range(200)) + "."
    blocks = [ParsedBlock(0, long, page_number=1, section_path="X",
                          metadata={"source_subtype": "pdf"})]
    chunker = DocumentChunker(target_tokens=100, overlap_tokens=15)
    chunks = chunker.chunk(blocks)
    assert len(chunks) >= 3, f"expected >=3 chunks for oversize block, got {len(chunks)}"
    for c in chunks:
        # Each chunk should be at most (target + some token-decode fuzz).
        assert c.token_count <= 100 + 40, f"chunk {c.chunk_index} = {c.token_count} tokens"


# ---------------------------------------------------------------------------
# Driver
# ---------------------------------------------------------------------------

def main() -> int:
    with section("4B - parsers"):
        check("4B", "pdf parser: 3 pages -> 3 blocks", test_pdf_parser_pages)
        check("4B", "docx parser: heading lineage + table flattening",
              test_docx_parser_section_path)
        check("4B", "xlsx parser: header attach + col_A fallback",
              test_xlsx_parser_header_attach_and_fallback)
        check("4B", "dispatcher: mime then extension; unsupported raises",
              test_dispatcher_mime_and_extension_fallback)

    with section("4B - DocumentChunker"):
        check("4B", "respects target_tokens budget", test_chunker_respects_target_tokens)
        check("4B", "overlap shingle survives across chunks",
              test_chunker_overlap_preserves_continuity)
        check("4B", "dominant page metadata inheritance",
              test_chunker_dominant_page_inheritance)
        check("4B", "empty + whitespace blocks are skipped",
              test_chunker_empty_and_whitespace_blocks_are_skipped)
        check("4B", "oversize single block sentence/hard-window split",
              test_chunker_oversize_single_block_falls_back)

    print("\n=== Summary ===")
    n_pass = sum(1 for r in results if r[2] == "PASS")
    n_fail = sum(1 for r in results if r[2] != "PASS")
    print(f"PASS: {n_pass}   FAIL: {n_fail}   TOTAL: {len(results)}")
    return 1 if n_fail else 0


if __name__ == "__main__":
    sys.exit(main())
