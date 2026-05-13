"""DOCX parser — python-docx with heading-aware section paths.

Strategy:

  - Walk the body in document order.
  - Maintain a `section_stack` of (level, text). Heading 1 sets level 1
    and clears anything deeper; Heading 2 nests under it; etc.
  - Every paragraph (or table row) becomes one ParsedBlock with
    `section_path = "/".join(stack texts)`. This way the chunker can
    surface "you said X in section 'Quarterly Review > North America'"
    at retrieval time.
  - Tables are flattened row-by-row as `"col1 | col2 | col3"` so the
    cell adjacency survives embedding.

We intentionally collapse very small blocks (< 8 tokens worth of words)
into the following block at the chunker level, *not* here — keeping the
parser block-faithful makes the chunker tests deterministic.
"""
from __future__ import annotations

import io
import logging
import re

from app.parsers.base import ParsedBlock

logger = logging.getLogger(__name__)

_WS_RX = re.compile(r"\s+")
_HEADING_RX = re.compile(r"^heading\s+(\d+)$", re.IGNORECASE)


def _heading_level(style_name: str | None) -> int | None:
    if not style_name:
        return None
    m = _HEADING_RX.match(style_name.strip())
    if not m:
        return None
    try:
        level = int(m.group(1))
    except ValueError:
        return None
    if 1 <= level <= 9:
        return level
    return None


def _flatten_section_stack(stack: list[tuple[int, str]]) -> str | None:
    if not stack:
        return None
    return " / ".join(text for _, text in stack)


def parse_docx(raw_bytes: bytes) -> list[ParsedBlock]:
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError

    try:
        document = Document(io.BytesIO(raw_bytes))
    except PackageNotFoundError as e:
        logger.warning("python-docx rejected file (%s) — returning empty blocks", e)
        return []
    except Exception as e:  # noqa: BLE001
        logger.warning("python-docx unexpected error (%s)", e)
        return []

    blocks: list[ParsedBlock] = []
    section_stack: list[tuple[int, str]] = []

    # python-docx exposes paragraphs and tables as separate ordered lists
    # but they share the body's XML order via `body.iter_inner_content()`
    # (1.0+) — we walk that to preserve "paragraph, table, paragraph"
    # ordering as it appears in the doc.
    try:
        body_iter = list(document.iter_inner_content())
    except AttributeError:
        # Older python-docx releases: fall back to paragraphs-then-tables.
        body_iter = list(document.paragraphs) + list(document.tables)

    for elem in body_iter:
        if hasattr(elem, "style") and hasattr(elem, "text"):
            # paragraph
            style_name = getattr(elem.style, "name", None)
            level = _heading_level(style_name)
            text = _WS_RX.sub(" ", (elem.text or "").strip()).strip()
            if level is not None:
                # Update section stack: pop anything at the same or deeper level.
                while section_stack and section_stack[-1][0] >= level:
                    section_stack.pop()
                if text:
                    section_stack.append((level, text))
                # We *also* emit the heading as a block — losing the
                # heading text loses search signal ("what's the title of
                # the section about pricing?").
                if text:
                    blocks.append(
                        ParsedBlock(
                            block_index=len(blocks),
                            text=text,
                            page_number=None,
                            section_path=_flatten_section_stack(section_stack[:-1]) or None,
                            metadata={
                                "source_subtype": "docx",
                                "block_role": "heading",
                                "heading_level": level,
                            },
                        )
                    )
                continue
            if not text:
                continue
            blocks.append(
                ParsedBlock(
                    block_index=len(blocks),
                    text=text,
                    page_number=None,
                    section_path=_flatten_section_stack(section_stack),
                    metadata={"source_subtype": "docx", "block_role": "paragraph"},
                )
            )
        elif hasattr(elem, "rows"):
            # table
            for row_idx, row in enumerate(elem.rows):
                cells = []
                for cell in row.cells:
                    cell_text = _WS_RX.sub(" ", (cell.text or "").strip()).strip()
                    if cell_text:
                        cells.append(cell_text)
                if not cells:
                    continue
                blocks.append(
                    ParsedBlock(
                        block_index=len(blocks),
                        text=" | ".join(cells),
                        page_number=None,
                        section_path=_flatten_section_stack(section_stack),
                        metadata={
                            "source_subtype": "docx",
                            "block_role": "table_row",
                            "row_index": row_idx,
                        },
                    )
                )
    return blocks
